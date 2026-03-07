#!/usr/bin/env python3
"""
Generate IAJIT-formatted Word document for Malak paper.
Format based on published IAJIT papers (Vol 23, No 1, 2026).
Two-column layout simulated via single-column Word (IAJIT handles typesetting).
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_font(run, name="Times New Roman", size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_iajit(doc, text, level=1):
    """IAJIT uses numbered headings: '1. Introduction' style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=11 if level == 1 else 10, bold=True)
    return p

def add_body(doc, text, indent_first=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_font(r, size=10, bold=True)
    r = p.add_run(text)
    set_font(r, size=10)
    return p

def add_numbered_item(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(f"{number}. ")
    set_font(r, size=10)
    r = p.add_run(text)
    set_font(r, size=10)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(code_text)
    set_font(run, name="Courier New", size=9)
    return p

def add_table(doc, caption, headers, rows):
    # Caption
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_font(r, size=9, bold=False, italic=False)

    # Table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=9, bold=True)
        # Gray background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'
        })
        shading.append(shd)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_font(r, size=9)

    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    return table


def main():
    doc = Document()

    # Page margins (IAJIT standard)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── Header line ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("The International Arab Journal of Information Technology")
    set_font(r, size=9, italic=True)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Malak: A Python Toolkit for Edge AI Model Optimization")
    set_font(r, size=16, bold=True)

    # ── Author (single author) ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Mohammed Alnemari")
    set_font(r, size=11, bold=True)
    r = p.add_run("\u00b9\u00b7\u00b2")
    set_font(r, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("\u00b9Department of Computer Engineering, Faculty of Computer Science and Information Technology, University of Tabuk, Tabuk, Saudi Arabia")
    set_font(r, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("\u00b2AIST Research Center, University of Tabuk, Tabuk, Saudi Arabia")
    set_font(r, size=9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("mnemari@ut.edu.sa")
    set_font(r, size=9, italic=True)

    # ── Abstract ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Abstract: ")
    set_font(r, size=10, bold=True)
    r = p.add_run(
        "Deploying deep learning models on resource-constrained edge devices requires systematic model compression and optimization. "
        "However, PyTorch's low-level quantization and pruning APIs require significant boilerplate code, making it difficult for researchers "
        "and students to quickly evaluate compression techniques. We present Malak, an open-source Python toolkit that provides a unified, "
        "high-level API for training, quantization (dynamic and static post-training quantization, quantization-aware training), pruning "
        "(magnitude and structured), knowledge distillation, ONNX model export, per-layer latency profiling, and distribution drift detection. "
        "Malak wraps PyTorch's quantization framework into task-oriented abstractions and exposes both a Python API and a command-line interface. "
        "We validate Malak on CIFAR-10 and Fashion-MNIST benchmarks using MobileNetV2, ResNet18, EfficientNet-B0, and a custom SimpleCNN. "
        "Quantization-aware training achieves 3.10x model compression with only 1.09% accuracy degradation. Dynamic post-training quantization "
        "preserves accuracy within 0.13% across all tested architectures. Magnitude pruning at 50% sparsity retains accuracy within 2.15% of the "
        "baseline after fine-tuning. The toolkit is available under the MIT license at https://github.com/alnemari-m/malak-platform."
    )
    set_font(r, size=10)

    # ── Keywords ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Keywords: ")
    set_font(r, size=10, bold=True)
    r = p.add_run("Edge AI, model compression, quantization, pruning, knowledge distillation, TinyML.")
    set_font(r, size=10, italic=True)

    # ── Received / Accepted ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Received; accepted")
    set_font(r, size=9, italic=True)

    # ══════════════════════════════════════════════════════
    # 1. Introduction
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "1. Introduction")

    add_body(doc, (
        "Edge AI\u2014deploying trained neural networks on microcontrollers, mobile phones, and embedded systems\u2014"
        "has become essential for latency-sensitive, privacy-preserving, and bandwidth-constrained applications [1, 2]. "
        "The global edge AI market is projected to grow significantly as more inference workloads move from cloud to device [2]. "
        "However, the gap between training a model in PyTorch and deploying it on a resource-constrained device remains wide. "
        "Practitioners must manually orchestrate quantization, pruning, model export, and validation steps, often writing ad-hoc scripts for each project."
    ))

    add_body(doc, (
        "PyTorch provides built-in quantization [3] and pruning APIs, but these are low-level: users must manually configure "
        "quantization configurations (qconfigs), insert observers, run calibration passes, and invoke conversion functions. "
        "A typical quantization-aware training (QAT) workflow requires approximately 40\u201360 lines of boilerplate code for "
        "observer insertion, training loop modification, and model conversion. This complexity is a barrier for researchers who "
        "want to quickly compare compression techniques and for students learning about edge deployment."
    ))

    add_body(doc, (
        "Several tools address parts of the model compression pipeline. TensorFlow Lite [5] and TFLite Micro [6] provide a "
        "complete path from TensorFlow to microcontrollers but require leaving the PyTorch ecosystem. Apache TVM [7] targets "
        "compiler-level graph optimizations but has a steep learning curve. Intel Neural Compressor [8] provides a comprehensive "
        "quantization framework but is oriented toward Intel hardware. CMSIS-NN [9] provides optimized inference kernels for ARM "
        "Cortex-M but assumes a pre-quantized model."
    ))

    add_body(doc, (
        "In this paper, we present Malak, an open-source Python toolkit that wraps PyTorch's quantization, pruning, and training "
        "APIs into a higher-level, task-oriented interface. Malak does not replace existing low-level frameworks; rather, it provides "
        "a structured starting point for evaluating and comparing model compression techniques within the PyTorch ecosystem."
    ))

    add_body(doc, "The primary contributions of this work are:", indent_first=True)
    add_numbered_item(doc, 1, "A unified Python API covering training, quantization (dynamic PTQ, static PTQ, QAT), pruning (magnitude and structured), knowledge distillation, ONNX export, per-layer profiling, and drift detection.")
    add_numbered_item(doc, 2, "A command-line interface (edgeai) for common model optimization tasks.")
    add_numbered_item(doc, 3, "Reproducible benchmark experiments on CIFAR-10 and Fashion-MNIST with four architectures and JSON result output.")
    add_numbered_item(doc, 4, "A proof-of-concept C inference implementation for ARM Cortex-M7 embedded targets.")

    add_body(doc, (
        "The remainder of this paper is organized as follows. Section 2 reviews related work. Section 3 describes the software "
        "architecture and key functionalities. Section 4 presents experimental validation. Section 5 discusses the results and "
        "limitations. Section 6 concludes the paper."
    ))

    # ══════════════════════════════════════════════════════
    # 2. Related Work
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "2. Related Work")

    add_heading_iajit(doc, "2.1. Model Quantization", level=2)
    add_body(doc, (
        "Post-training quantization (PTQ) reduces model precision after training without retraining. Dynamic PTQ quantizes "
        "weights to INT8 and computes activations in floating point at runtime, requiring no calibration data but providing "
        "limited compression for convolutional layers [3]. Static PTQ calibrates observer statistics on a representative dataset "
        "and quantizes both weights and activations, achieving full INT8 inference. Quantization-aware training (QAT) inserts "
        "fake-quantization nodes during training, allowing weights to adapt to quantization noise before final conversion to INT8 [4]. "
        "PyTorch implements all three approaches through its torch.quantization namespace."
    ))

    add_heading_iajit(doc, "2.2. Model Pruning", level=2)
    add_body(doc, (
        "Pruning removes redundant weights or structures from neural networks [10]. Unstructured pruning zeros individual weights "
        "based on magnitude [11], producing sparse weight matrices. Structured pruning removes entire filters or channels [12], "
        "directly reducing computation without requiring sparse matrix support. Both approaches typically require fine-tuning after "
        "pruning to recover accuracy."
    ))

    add_heading_iajit(doc, "2.3. Knowledge Distillation", level=2)
    add_body(doc, (
        "Knowledge distillation transfers knowledge from a large teacher model to a smaller student model using temperature-scaled "
        "soft targets [13]. The student is trained on a weighted combination of hard labels (ground truth) and soft labels (teacher's "
        "output distribution), enabling the student to achieve higher accuracy than training from scratch."
    ))

    add_heading_iajit(doc, "2.4. Existing Toolkits", level=2)
    add_body(doc, (
        "Table 1 compares Malak with existing model optimization tools. TensorFlow Lite [5] provides comprehensive quantization "
        "and deployment but is tied to the TensorFlow ecosystem. Intel Neural Compressor [8] offers extensive quantization support "
        "for PyTorch but is optimized for Intel hardware (CPUs, Habana accelerators). Apache TVM [7] provides compiler-level "
        "optimizations across hardware backends but requires understanding of compilation graphs. CMSIS-NN [9] provides hand-optimized "
        "inference kernels for ARM Cortex-M processors. Malak operates at a higher abstraction level than these tools, providing "
        "task-oriented wrappers around PyTorch's built-in compression APIs with minimal setup requirements."
    ))

    # Table 1: Comparison
    add_table(doc,
        "Table 1. Comparison of model optimization tools.",
        ["Feature", "Malak", "TFLite", "INC", "TVM", "CMSIS"],
        [
            ["PyTorch native",  "\u2713", "",     "\u2713", "\u2713", ""],
            ["Dynamic PTQ",     "\u2713", "\u2713", "\u2713", "",     ""],
            ["Static PTQ",      "\u2713", "\u2713", "\u2713", "",     ""],
            ["QAT",             "\u2713", "\u2713", "\u2713", "",     ""],
            ["Pruning",         "\u2713", "",     "\u2713", "",       ""],
            ["Distillation",    "\u2713", "",     "\u2713", "",       ""],
            ["ONNX export",     "\u2713", "",     "\u2713", "\u2713", ""],
            ["CLI interface",   "\u2713", "\u2713", "\u2713", "\u2713", ""],
            ["Edge kernels",    "",       "\u2713", "",     "\u2713", "\u2713"],
            ["Graph compiler",  "",       "\u2713", "",     "\u2713", ""],
            ["Setup (LoC)",     "~5",     "~15",  "~20",   "~30",   "N/A"],
        ]
    )

    # ══════════════════════════════════════════════════════
    # 3. Software Architecture
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "3. Software Architecture")

    add_body(doc, (
        "Malak is organized into six modules (Figure 1), installed via pip install -e . and usable through both a Python API "
        "and a CLI. The modules are grouped into three layers: model preparation (Training, Quantization, Compression), deployment "
        "and analysis (Compiler, Runtime, Monitoring), and the command-line interface."
    ))

    add_body(doc, "[Figure 1: Malak software architecture diagram \u2014 insert Malak_Architecture.svg here]", indent_first=False)

    add_heading_iajit(doc, "3.1. Training Module", level=2)
    add_body(doc, (
        "The Trainer class provides a configurable training loop supporting SGD and Adam optimizers with cosine annealing scheduling. "
        "Dataset loaders for CIFAR-10 and Fashion-MNIST include standard data augmentation (random crop, horizontal flip) and "
        "normalization. The trainer automatically tracks the best model checkpoint and reports per-epoch accuracy and latency."
    ))

    add_heading_iajit(doc, "3.2. Quantization Module", level=2)
    add_body(doc, "The quantization module provides three classes:")
    add_bullet(doc, " Wraps torch.quantization.quantize_dynamic to quantize Linear layers to INT8. Requires no calibration data. Suitable for models dominated by fully-connected layers.", bold_prefix="\u2022 DynamicPTQ:")
    add_bullet(doc, " Inserts observers, runs a calibration pass over a representative dataset, and converts all supported layers to INT8. Produces full model compression.", bold_prefix="\u2022 StaticPTQ:")
    add_bullet(doc, " A three-step API (prepare \u2192 train \u2192 convert) that inserts fake-quantization observers during training, allowing weights to adapt to quantization noise before final INT8 conversion.", bold_prefix="\u2022 QAT:")

    add_body(doc, "The QAT workflow is illustrated below:")
    add_code_block(doc, (
        "from malak.quantization import QAT\n"
        "qat = QAT()\n"
        "model = qat.prepare(model, backend=\"fbgemm\")\n"
        "model = qat.train(model, train_loader,\n"
        "                  test_loader, epochs=5)\n"
        "quantized = qat.convert(model)"
    ))

    add_heading_iajit(doc, "3.3. Compression Module", level=2)
    add_body(doc, (
        "The compression module provides magnitude pruning (MagnitudePruner), structured filter pruning (StructuredPruner), "
        "and knowledge distillation (KnowledgeDistiller). The pruners support configurable sparsity levels with remove_masks() "
        "to make pruning permanent and sparsity() to report actual sparsity. The distiller trains a student model using a weighted "
        "combination of cross-entropy (hard labels) and KL-divergence (temperature-scaled soft labels from the teacher), following [13]."
    ))

    add_heading_iajit(doc, "3.4. Supporting Modules", level=2)
    add_body(doc, (
        "The Compiler module exports models to ONNX format with optional validation. The Runtime module benchmarks inference "
        "latency and throughput. The Monitoring module provides per-layer latency profiling via forward hooks and KL-divergence-based "
        "distribution drift detection for deployed models. The edgeai CLI exposes all modules through subcommands (train, quantize, "
        "evaluate, export, profile)."
    ))

    # ══════════════════════════════════════════════════════
    # 4. Experimental Validation
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "4. Experimental Validation")

    add_body(doc, (
        "We validate Malak through four experiments covering quantization, architecture comparison, pruning, and cross-dataset "
        "evaluation. All experiments use CPU inference (Intel Core i7) with PyTorch 2.x. Results are stored as JSON files in the "
        "repository and every number reported is directly measured. All experiments are reproducible via scripts in the experiments/ directory."
    ))

    add_heading_iajit(doc, "4.1. Experiment 1: MobileNetV2 Quantization", level=2)
    add_body(doc, (
        "We train MobileNetV2 [14] (2.24M parameters, first convolution modified for 32\u00d732 input) on CIFAR-10 for 15 epochs "
        "with SGD (lr=0.01, momentum=0.9, cosine annealing). We then apply dynamic PTQ and QAT (5 fine-tuning epochs with lr=0.001). "
        "Table 2 shows the results."
    ))

    add_table(doc,
        "Table 2. MobileNetV2 quantization results on CIFAR-10.",
        ["Method", "Acc. (%)", "Size (MB)", "Compr."],
        [
            ["FP32 Baseline", "77.81", "8.76", "1.00\u00d7"],
            ["Dynamic PTQ",   "77.84", "8.73", "1.00\u00d7"],
            ["QAT (INT8)",    "76.72", "2.82", "3.10\u00d7"],
        ]
    )

    add_body(doc, (
        "Dynamic PTQ preserves accuracy (within 0.03%) but provides minimal size reduction because it only quantizes Linear layers "
        "while the convolutional layers\u2014which dominate MobileNetV2\u2014remain in FP32. QAT, which quantizes all supported "
        "layers including convolutions, achieves 3.10\u00d7 compression with 1.09% accuracy drop."
    ))

    add_body(doc, (
        "We note that the 77.81% FP32 baseline is below published state-of-the-art for MobileNetV2 on CIFAR-10 (~94%) due to our "
        "shortened training schedule (15 epochs vs. typical 200+ epochs with advanced augmentation). The purpose of this experiment "
        "is to validate the toolkit\u2019s quantization pipeline, not to achieve maximum accuracy. The relative compression ratios "
        "and accuracy drops are consistent with published quantization studies [3, 4]."
    ))

    add_heading_iajit(doc, "4.2. Experiment 2: Architecture Comparison", level=2)
    add_body(doc, (
        "We compare ResNet18 (11.2M parameters) and EfficientNet-B0 [15] (4.0M parameters) on CIFAR-10, training each for 15 epochs "
        "with SGD and applying dynamic INT8 PTQ. Table 3 shows the results."
    ))

    add_table(doc,
        "Table 3. Architecture comparison with dynamic INT8 PTQ on CIFAR-10.",
        ["Model", "FP32", "INT8", "Drop", "Compr."],
        [
            ["ResNet18",       "87.41%", "87.38%", "0.03%", "1.00\u00d7"],
            ["EfficientNet-B0", "66.81%", "66.68%", "0.13%", "1.00\u00d7"],
        ]
    )

    add_body(doc, (
        "Dynamic PTQ preserves accuracy within 0.13% for both architectures. The 1.00\u00d7 compression ratio is expected: dynamic "
        "PTQ only quantizes nn.Linear layers at runtime, and both architectures are dominated by convolutional layers whose weights "
        "remain in FP32. The EfficientNet-B0 baseline accuracy (66.81%) is lower than ResNet18 (87.41%) because EfficientNet's "
        "compound scaling was designed for 224\u00d7224 ImageNet inputs and requires longer training schedules to converge on "
        "32\u00d732 CIFAR-10 images."
    ))

    add_heading_iajit(doc, "4.3. Experiment 3: Pruning", level=2)
    add_body(doc, (
        "We apply L1-unstructured magnitude pruning at 30%, 50%, and 70% target sparsity to a trained MobileNetV2, followed by "
        "3 epochs of fine-tuning per sparsity level. We also test L2-structured filter pruning at 30% and 50% sparsity. "
        "Table 4 shows the results."
    ))

    add_table(doc,
        "Table 4. Pruning results on CIFAR-10 MobileNetV2 (with fine-tuning).",
        ["Method", "Target", "Acc. (%)", "Drop", "Actual"],
        [
            ["Baseline",    "0%",  "77.81", "\u2014",   "0%"],
            ["Magnitude",   "30%", "77.32", "0.49%",  "29.5%"],
            ["Magnitude",   "50%", "75.66", "2.15%",  "49.2%"],
            ["Magnitude",   "70%", "70.47", "7.34%",  "68.9%"],
            ["Structured",  "30%", "72.86", "4.95%",  "29.6%"],
            ["Structured",  "50%", "64.17", "13.64%", "49.2%"],
        ]
    )

    add_body(doc, (
        "Magnitude pruning at 30% sparsity causes negligible accuracy loss (0.49%), while 50% sparsity results in a moderate 2.15% "
        "drop. At 70% sparsity, accuracy drops by 7.34%, indicating that more aggressive fine-tuning (e.g., 10\u201320 epochs) would "
        "be beneficial at high sparsity levels. Structured pruning causes larger accuracy drops than magnitude pruning at the same "
        "sparsity level because removing entire filters is more destructive than zeroing individual weights. This is consistent with "
        "findings in the pruning literature [11, 12]."
    ))

    add_heading_iajit(doc, "4.4. Experiment 4: Fashion-MNIST with SimpleCNN", level=2)
    add_body(doc, (
        "We train a custom SimpleCNN (458K parameters: three convolutional layers with max-pooling, followed by two fully-connected "
        "layers) on Fashion-MNIST for 20 epochs with Adam (lr=0.001) and apply dynamic INT8 PTQ. Table 5 shows the results."
    ))

    add_table(doc,
        "Table 5. SimpleCNN results on Fashion-MNIST with dynamic PTQ.",
        ["Method", "Acc. (%)", "Size (MB)", "Compr."],
        [
            ["FP32 Baseline", "92.19", "1.75", "1.00\u00d7"],
            ["Dynamic PTQ",   "92.20", "0.60", "2.91\u00d7"],
        ]
    )

    add_body(doc, (
        "Unlike the convolutional-heavy architectures in Experiments 1\u20132, SimpleCNN has substantial fully-connected layers "
        "(64\u00d77\u00d77 \u2192 128 \u2192 10), which are quantized by dynamic PTQ. This results in 2.91\u00d7 compression with no "
        "accuracy loss, demonstrating that dynamic PTQ is effective for FC-heavy models. The 92.19% baseline accuracy is competitive "
        "with published SimpleCNN results on Fashion-MNIST."
    ))

    add_heading_iajit(doc, "4.5. Embedded Proof-of-Concept", level=2)
    add_body(doc, (
        "The repository includes a C implementation of SimpleCNN inference using INT8 weights exported by a Python script. "
        "The C code implements naive conv2d, maxpool, fully-connected, and ReLU operations with symmetric INT8 quantization, "
        "targeting ARM Cortex-M7 (STM32H7) via arm-none-eabi-gcc. This implementation serves as a proof-of-concept for the export "
        "pipeline; production-quality embedded inference would require optimized kernels from CMSIS-NN [9] or TFLite Micro [6]."
    ))

    # ══════════════════════════════════════════════════════
    # 5. Discussion
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "5. Discussion")

    add_heading_iajit(doc, "5.1. Key Findings", level=2)
    add_body(doc, (
        "The experimental results demonstrate three findings: (1) QAT is the most effective quantization method for convolutional "
        "models, achieving 3.10\u00d7 compression with minimal accuracy loss, while dynamic PTQ provides negligible compression for "
        "conv-heavy architectures but is effective for FC-heavy models (2.91\u00d7 on SimpleCNN). (2) Magnitude pruning is more "
        "accuracy-preserving than structured pruning at the same sparsity level, at the cost of requiring sparse matrix support for "
        "inference speedup. (3) The Malak API successfully produces reproducible, measurable results across four architectures and "
        "two datasets, validating its correctness."
    ))

    add_heading_iajit(doc, "5.2. Limitations", level=2)
    add_body(doc, "This work has several limitations that should be acknowledged:")
    add_numbered_item(doc, 1, "Limited baselines. Our FP32 accuracies (77.81% MobileNetV2, 66.81% EfficientNet-B0) are below published state-of-the-art due to abbreviated training schedules. While the relative compression metrics remain valid, results at higher baselines may differ.")
    add_numbered_item(doc, 2, "Single-run results. All experiments report single-run results without standard deviations. Stochastic training introduces variance that is not captured.")
    add_numbered_item(doc, 3, "CPU-only evaluation. All latency measurements are on desktop CPU, not on actual edge hardware. Inference speedup from quantization depends heavily on hardware support for INT8 operations.")
    add_numbered_item(doc, 4, "Limited scope. The toolkit supports only CIFAR-10 and Fashion-MNIST datasets and a small set of architectures. Extension to larger datasets (ImageNet) and domain-specific tasks is left to users.")
    add_numbered_item(doc, 5, "Wrapper-level contribution. Malak wraps existing PyTorch APIs rather than introducing new algorithms. Its value lies in usability and reproducibility, not algorithmic novelty.")
    add_numbered_item(doc, 6, "Embedded deployment. The C inference code is a proof-of-concept with floating-point accumulation, not a production-ready embedded runtime.")

    add_heading_iajit(doc, "5.3. Comparison with Existing Tools", level=2)
    add_body(doc, (
        "Malak's primary advantage over raw PyTorch is reduced boilerplate: the QAT workflow requires 5 lines of code (the code "
        "example in Section 3) compared to ~40 lines using PyTorch directly. Compared to Intel Neural Compressor [8], Malak is simpler "
        "to install (pure Python, no hardware-specific dependencies) but offers fewer quantization strategies and no hardware-aware "
        "optimization. Compared to TFLite, Malak stays within the PyTorch ecosystem but lacks a production-quality embedded runtime. "
        "Malak is best suited for rapid prototyping and educational use rather than production deployment."
    ))

    # ══════════════════════════════════════════════════════
    # 6. Conclusions
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "6. Conclusions")

    add_body(doc, (
        "We presented Malak, an open-source Python toolkit for edge AI model optimization that provides a unified API for "
        "quantization, pruning, distillation, ONNX export, profiling, and drift detection. Validation across four experiments on "
        "CIFAR-10 and Fashion-MNIST demonstrates that the toolkit produces correct and reproducible compression results. QAT achieves "
        "3.10\u00d7 compression with 1.09% accuracy drop on MobileNetV2, while magnitude pruning at 50% sparsity retains accuracy "
        "within 2.15% of the baseline."
    ))

    add_body(doc, (
        "Future work includes: (1) integration with CMSIS-NN for production-quality embedded inference, (2) support for mixed-precision "
        "quantization [16], (3) multi-run experiments with statistical reporting, (4) extending the training pipeline to ImageNet-scale "
        "datasets, and (5) automated compression policy search combining quantization and pruning."
    ))

    add_body(doc, "The toolkit is available at https://github.com/alnemari-m/malak-platform under the MIT license.")

    # ── Acknowledgments ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Acknowledgments")
    set_font(r, size=11, bold=True)

    add_body(doc, (
        "The authors thank the open-source PyTorch community for the quantization and pruning APIs upon which Malak is built."
    ))

    # ══════════════════════════════════════════════════════
    # References
    # ══════════════════════════════════════════════════════
    add_heading_iajit(doc, "References")

    refs = [
        '[1] Warden P. and Situnayake D., TinyML: Machine Learning with TensorFlow Lite on Arduino and Ultra-Low-Power Microcontrollers, O\'Reilly Media, 2019.',
        '[2] Zhou Z., Chen X., Li E., Zeng L., Luo K., and Zhang J., "Edge Intelligence: The Confluence of Edge Computing and Artificial Intelligence," IEEE Internet of Things Journal, vol. 7, no. 8, pp. 7457-7469, 2020.',
        '[3] Jacob B. et al., "Quantization and Training of Neural Networks for Efficient Integer-Arithmetic-Only Inference," in Proceedings of IEEE CVPR, pp. 2704-2713, 2018.',
        '[4] Bhalgat Y., Lee J., Nagel M., Blankevoort T., and Kwak N., "Quantization Aware Training for Low Bit-Width Neural Networks," in NeurIPS Workshop, 2020.',
        '[5] Abadi M. et al., "TensorFlow Lite: On-Device Machine Learning Framework," Google Research, 2019.',
        '[6] David R. et al., "TensorFlow Lite Micro: Embedded Machine Learning for TinyML Systems," in Proceedings of MLSys, vol. 3, pp. 800-811, 2021.',
        '[7] Chen T. et al., "TVM: An Automated End-to-End Optimizing Compiler for Deep Learning," in Proceedings of USENIX OSDI, pp. 578-594, 2018.',
        '[8] Intel Corporation, "Intel Neural Compressor: An Open-Source Python Library for Model Compression," GitHub Repository, 2021.',
        '[9] Lai L., Suda N., and Chandra V., "CMSIS-NN: Efficient Neural Network Kernels for Arm Cortex-M CPUs," arXiv preprint arXiv:1801.06601, 2018.',
        '[10] Deng L., Li G., Han S., Shi L., and Xie Y., "Model Compression and Hardware Acceleration for Neural Networks: A Comprehensive Survey," Proceedings of IEEE, vol. 108, no. 4, pp. 485-532, 2020.',
        '[11] Frankle J. and Carlin M., "The Lottery Ticket Hypothesis: Finding Sparse, Trainable Neural Networks," arXiv preprint arXiv:1803.03635, 2018.',
        '[12] Liu Z. et al., "Learning Efficient Convolutional Networks Through Network Slimming," in Proceedings of IEEE ICCV, pp. 2736-2744, 2017.',
        '[13] Hinton G., Vinyals O., and Dean J., "Distilling the Knowledge in a Neural Network," in NIPS Deep Learning Workshop, 2015.',
        '[14] Howard A. et al., "Searching for MobileNetV3," in Proceedings of IEEE ICCV, pp. 1314-1324, 2019.',
        '[15] Tan M. and Le Q., "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks," in Proceedings of ICML, pp. 6105-6114, 2019.',
        '[16] Wang K., Liu Z., Lin Y., Lin J., and Han S., "HAQ: Hardware-Aware Automated Quantization with Mixed Precision," in Proceedings of IEEE CVPR, pp. 8612-8620, 2019.',
        '[17] Bai J., Lu F., Zhang K. et al., "ONNX: Open Neural Network Exchange," GitHub Repository, 2019.',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        r = p.add_run(ref)
        set_font(r, size=9)

    # ── Author bios (IAJIT includes short bios) ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)

    for name, bio in [
        ("Mohammed Hassan Alnemari", "received the B.Sc. degree in computer engineering from Taif University, Saudi Arabia, and the M.S. and Ph.D. degrees in computer engineering from the University of California, Irvine (UCI), USA. He is currently an Assistant Professor and Chair of the Computer Engineering Department at the University of Tabuk, Saudi Arabia. His research interests include Edge Artificial Intelligence (Edge AI), TinyML, neural network compression, efficient deep learning architectures, and hardware\u2013software co-design for intelligent embedded systems. His work focuses on enabling energy-efficient machine learning inference on resource-constrained edge devices."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(name + " ")
        set_font(r, size=9, bold=True)
        r = p.add_run(bio)
        set_font(r, size=9)

    # Save
    out = os.path.expanduser("~/Desktop/Malak_IAJIT_Final.docx")
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
